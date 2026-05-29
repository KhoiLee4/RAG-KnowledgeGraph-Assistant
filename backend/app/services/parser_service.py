"""
parser_service.py — Dịch vụ extract văn bản từ các định dạng file khác nhau.

Hỗ trợ:
  - PDF       → pypdf
  - DOCX      → python-docx
  - XLSX/CSV  → openpyxl / csv
  - Ảnh       → Tesseract OCR (optional/bonus)
  - Text      → decode UTF-8

Dùng parse_file(file_bytes, mime_type) để tự động chọn parser phù hợp.
"""

import csv
import io
import logging
import os
from typing import Any

from app.core.config import settings

logger = logging.getLogger(__name__)

# Ánh xạ MIME type → method parse
MIME_TO_PARSER: dict[str, str] = {
    "application/pdf": "parse_pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "parse_docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "parse_xlsx",
    "text/plain": "parse_text",
    "text/markdown": "parse_text",
    "text/html": "parse_text",
    "text/csv": "parse_csv",
    "image/jpeg": "parse_image_ocr",
    "image/png": "parse_image_ocr",
    "image/webp": "parse_image_ocr",
    "image/bmp": "parse_image_ocr",
}


class ParserService:
    """
    Router tự động chọn parser dựa trên MIME type của file.
    Tất cả method trả về chuỗi văn bản thuần (plain text).
    """

    # ── PDF ───────────────────────────────────────────────────

    def parse_pdf(self, file_bytes: bytes) -> str:
        """
        Trích xuất văn bản từ file PDF.
        Mỗi trang được đánh dấu [Trang N] để dễ theo dõi vị trí.

        Args:
            file_bytes: Nội dung file PDF dạng bytes.

        Returns:
            Chuỗi văn bản từ tất cả trang PDF.
        """
        try:
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages: list[str] = []

            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Trang {i}]\n{text.strip()}")

            result = "\n\n".join(pages)
            logger.info("PDF: extract %d ký tự từ %d trang.", len(result), len(reader.pages))
            return result

        except ImportError:
            raise ImportError("Cài pypdf: pip install pypdf")
        except Exception as e:
            logger.error("parse_pdf lỗi: %s", e)
            raise

    # ── DOCX ──────────────────────────────────────────────────

    def parse_docx(self, file_bytes: bytes) -> str:
        """
        Trích xuất văn bản từ file Word (.docx).
        Giữ nguyên cấu trúc đoạn văn, bỏ qua các đoạn trống.

        Args:
            file_bytes: Nội dung file DOCX dạng bytes.

        Returns:
            Chuỗi văn bản từ tất cả đoạn trong tài liệu.
        """
        try:
            import docx

            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Trích thêm text từ bảng trong document
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            result = "\n\n".join(paragraphs)
            logger.info("DOCX: extract %d ký tự.", len(result))
            return result

        except ImportError:
            raise ImportError("Cài python-docx: pip install python-docx")
        except Exception as e:
            logger.error("parse_docx lỗi: %s", e)
            raise

    # ── XLSX ──────────────────────────────────────────────────

    def parse_xlsx(self, file_bytes: bytes) -> str:
        """
        Trích xuất dữ liệu từ file Excel (.xlsx).
        Mỗi sheet được đánh tiêu đề [Sheet: TÊN], mỗi hàng là một dòng.

        Args:
            file_bytes: Nội dung file XLSX dạng bytes.

        Returns:
            Chuỗi văn bản đại diện nội dung spreadsheet.
        """
        try:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            lines: list[str] = []

            for sheet in wb.worksheets:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    # Bỏ hàng hoàn toàn trống
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        lines.append("\t".join(cells))

            result = "\n".join(lines)
            logger.info("XLSX: extract %d ký tự từ %d sheet.", len(result), len(wb.worksheets))
            return result

        except ImportError:
            raise ImportError("Cài openpyxl: pip install openpyxl")
        except Exception as e:
            logger.error("parse_xlsx lỗi: %s", e)
            raise

    # ── CSV ───────────────────────────────────────────────────

    def parse_csv(self, file_bytes: bytes) -> str:
        """
        Trích xuất dữ liệu từ file CSV.

        Args:
            file_bytes: Nội dung file CSV dạng bytes.

        Returns:
            Chuỗi văn bản với mỗi hàng là một dòng.
        """
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text))
            rows = ["\t".join(row) for row in reader if any(row)]
            result = "\n".join(rows)
            logger.info("CSV: extract %d ký tự.", len(result))
            return result
        except Exception as e:
            logger.error("parse_csv lỗi: %s", e)
            raise

    # ── Text / Markdown ───────────────────────────────────────

    def parse_text(self, file_bytes: bytes) -> str:
        """
        Decode file văn bản thuần (txt, md, html) sang string UTF-8.

        Args:
            file_bytes: Nội dung file dạng bytes.

        Returns:
            Chuỗi văn bản.
        """
        result = file_bytes.decode("utf-8", errors="replace").strip()
        logger.info("Text: decode %d ký tự.", len(result))
        return result

    # ── Ảnh → OCR (Optional/Bonus) ───────────────────────────

    def parse_image_ocr(self, file_bytes: bytes) -> str:
        """
        [OPTIONAL] Trích xuất văn bản từ ảnh bằng Tesseract OCR.
        Cần cài Tesseract binary và pytesseract + Pillow.

        Args:
            file_bytes: Nội dung file ảnh dạng bytes (JPEG, PNG, WebP...).

        Returns:
            Văn bản trích xuất từ ảnh. Trả về "" nếu Tesseract chưa cài.
        """
        try:
            import pytesseract
            from PIL import Image, ImageEnhance, ImageFilter

            # Cấu hình đường dẫn Tesseract (cần thiết trên Windows)
            if os.name == "nt":
                pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

            img = Image.open(io.BytesIO(file_bytes))

            # Tiền xử lý ảnh để tăng độ chính xác OCR
            if img.mode != "L":
                img = img.convert("L")  # Grayscale
            img = ImageEnhance.Contrast(img).enhance(2.0)
            img = img.filter(ImageFilter.SHARPEN)

            text = pytesseract.image_to_string(img, lang=settings.OCR_LANG, config="--psm 3")
            result = text.strip()
            logger.info("OCR: extract %d ký tự từ ảnh.", len(result))
            return result

        except ImportError:
            logger.warning("Tesseract/pytesseract chưa cài — bỏ qua OCR.")
            return ""
        except Exception as e:
            logger.error("parse_image_ocr lỗi: %s", e)
            return ""

    # ── Router tự động ────────────────────────────────────────

    def parse_file(self, file_bytes: bytes, mime_type: str) -> str:
        """
        Tự động chọn parser phù hợp dựa trên MIME type của file.
        Đây là method chính cần gọi trong pipeline indexing.

        Args:
            file_bytes: Nội dung file dạng bytes.
            mime_type: MIME type của file, ví dụ "application/pdf".

        Returns:
            Văn bản thuần đã trích xuất.

        Raises:
            ValueError: Nếu MIME type chưa được hỗ trợ.
        """
        # Chuẩn hóa mime_type (bỏ tham số phụ như charset)
        base_mime = mime_type.split(";")[0].strip().lower()

        # Tìm parser method
        parser_name = MIME_TO_PARSER.get(base_mime)

        if parser_name:
            parser_method = getattr(self, parser_name)
            logger.info("Dùng parser '%s' cho MIME '%s'.", parser_name, base_mime)
            return parser_method(file_bytes)

        # Fallback: thử decode UTF-8 nếu không biết MIME type
        logger.warning(
            "MIME type '%s' chưa có parser, thử decode UTF-8.", base_mime
        )
        return file_bytes.decode("utf-8", errors="replace").strip()

    def get_supported_types(self) -> list[str]:
        """Trả về danh sách MIME type được hỗ trợ."""
        return list(MIME_TO_PARSER.keys())


# ── Test độc lập ──────────────────────────────────────────────
if __name__ == "__main__":
    import os

    print("=== Test ParserService ===")
    parser = ParserService()

    print("MIME type được hỗ trợ:")
    for t in parser.get_supported_types():
        print(f"  - {t}")

    # Test parse text
    sample_text = "Xin chào!\nĐây là file thử nghiệm.\nDòng thứ ba."
    result = parser.parse_file(sample_text.encode(), "text/plain")
    print(f"\nParse text:\n  Input: {sample_text!r}")
    print(f"  Output: {result!r}")

    # Test parse PDF (nếu có file test)
    test_pdf = "test_sample.pdf"
    if os.path.exists(test_pdf):
        with open(test_pdf, "rb") as f:
            pdf_bytes = f.read()
        pdf_text = parser.parse_pdf(pdf_bytes)
        print(f"\nParse PDF '{test_pdf}':\n  {pdf_text[:200]}...")
    else:
        print(f"\nKhông có '{test_pdf}' để test PDF parser.")
